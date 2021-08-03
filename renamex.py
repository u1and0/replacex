"""
docxファイルの文字列置換を行うモジュール
usage:
    $ python xreplace.py oldword newword filename
    $ python xreplace.py -f oldfile.docx -o newfile.docx -w replaceword.csv
"""
import sys
import csv
from docx import Document
# from docx.shared import Pt
from docx.shared import RGBColor


def replace_text(paragraph, before, after):
    """paragraph内の文字列beforeをafterへ置換する"""
    replaced_text = paragraph.text.replace(before, after)
    if paragraph.text != replaced_text:
        paragraph.text = replaced_text
        # paragraph.runs[0].font.size = Pt(12)
        paragraph.runs[0].font.color.rgb = RGBColor(235, 0, 0)


def replace_texts(readfile, writefile, *words):
    """
    readfile内の本文とテーブルの文字列を
    wordsに従って置換して
    writefileへ保存する。
    """
    document = Document(readfile)
    for word in words:
        # 本文書き換え
        for paragraph in document.paragraphs:
            replace_text(paragraph, word[0], word[1])
        # テーブル書き換え
        paragraphs = (paragraph for table in document.tables
                      for row in table.rows for cell in row.cells
                      for paragraph in cell.paragraphs)
        for paragraph in paragraphs:
            replace_text(paragraph, word[0], word[1])
    document.save(writefile)


if __name__ == '__main__':
    if len(sys.argv) < 4:
        raise ValueError('引数が足りません' + __doc__)
    old, new, filename = sys.argv[1], sys.argv[2], sys.argv[3]
    document = Document(filename)
    # 本文書き換え
    for paragraph in document.paragraphs:
        replace_text(paragraph, old, new)
    # テーブル書き換え
    paragraphs = (paragraph for table in document.tables for row in table.rows
                  for cell in row.cells for paragraph in cell.paragraphs)
    for paragraph in paragraphs:
        replace_text(paragraph, old, new)
    document.save(filename)
