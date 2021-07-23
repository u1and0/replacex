"""
途中でgolangでやりたくなったので放棄
[python-docxを使ってdocxファイルを書き換える](https://qiita.com/butada/items/93ad2301348a4fe28b5c)
"""
import sys
import csv
from docx import Document
# from docx.shared import Pt
from docx.shared import RGBColor


def replace_text(paragraph, before, after):
    """paragraph内の文字列beforeをafterへ置換する"""
    replaced_text = paragraph.text.replace(before, after)
    if paragraph.text != replaced_text:
        paragraph.text = replaced_text
        # paragraph.runs[0].font.size = Pt(12)
        paragraph.runs[0].font.color.rgb = RGBColor(204, 0, 0)


def replace_texts(readfile, writefile, *words):
    """
    readfile内の本文とテーブルの文字列を
    REPLACERに従って置換して
    writefileへ保存する。
    """
    document = Document(readfile)
    for word in words:
        # 本文書き換え
        for paragraph in document.paragraphs:
            replace_text(paragraph, *word)
        # テーブル書き換え
        paragraphs = (paragraph for table in document.tables
                      for row in table.rows for cell in row.cells
                      for paragraph in cell.paragraphs)
        for paragraph in paragraphs:
            replace_text(paragraph, *word)
    document.save(writefile)


if __name__ == '__main__':
    if len(sys.argv) < 4:
        """old code"""
        # HARD CORDED HERE
        READFILE = '/4_Restriction/SS用無線通信総合説明書/j情報セキュリティ（19SS以降）/27SS/帳票類/' +\
            '25　25SS可搬記憶媒体持ち込み申請書2.docx'
        # replace file name
        WRITEFILE = READFILE.replace('25SS', '27SS')
        # replace word list
        REPLACER = (
            # (BEFORE, AFTER)
            ('25SS', '27SS'),
            ('23SS', '25SS'),
        )
        replace_texts(READFILE, WRITEFILE, *REPLACER)
        sys.exit(1)
    # Command line here
    readfile, writefile = sys.argv[1], sys.argv[2]
    with open(sys.argv[3], newline='', encoding='sjis') as csvfile:
        reader = csv.reader(csvfile, delimiter=(','), quotechar='|')
        replacer = [row[0].split('\t') for row in reader]
    replace_texts(readfile, writefile, *replacer)
