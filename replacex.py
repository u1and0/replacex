#!/usr/bin/env python3
"""Replace text and highlight it in docx files.

# require: python-docx

$ pip install python-docx
"""
import sys
import argparse
from itertools import chain
from more_itertools import collapse
from docx import Document
from docx.shared import RGBColor

VERSION = 'v1.0.2'
CRED = '\033[91m'
CEND = '\033[0m'


def replace_text(paragraph, before, after):
    """replace before -> after in paragraph"""
    replaced_text = paragraph.text.replace(before, after)
    if paragraph.text != replaced_text:
        paragraph.text = replaced_text
        paragraph.runs[0].font.color.rgb = RGBColor(235, 0, 0)
        yield paragraph.text


def replace_document(old, new, document):
    """execute replaced_text to multiple paragraphs"""
    # Rewrite sentence
    sentence_paragraphs = (paragraph for paragraph in document.paragraphs)
    # Rewrite table
    table_paragraphs = (paragraph for table in document.tables
                        for row in table.rows for cell in row.cells
                        for paragraph in cell.paragraphs)
    # Concat iter
    paragraphs = chain(sentence_paragraphs, table_paragraphs)
    # Edit contents
    return (replace_text(paragraph, old, new) for paragraph in paragraphs)


def main(old, new, *filenames, dryrun, verbose):
    """execute replace_document to multiple files"""
    for filename in filenames:
        document = Document(filename)
        if verbose or dryrun:
            print("==filename:", filename, "==")
        paragraphs = replace_document(old, new, document)
        # Break nested iters by collapse()
        # replace_document() , replace_text() are generator
        # execute to edit docx through this 'for' statement
        for paragraph_text in collapse(paragraphs):  # Replace text HERE
            # Print out result to stdout if verbose mode or dryrun mode
            if verbose or dryrun:
                colored = paragraph_text.replace(new, CRED + new + CEND)
                print(colored)
        # Save Document unless dryrun mode
        if not dryrun:
            document.save(filename)


def parse():
    """arg parser"""
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument('old', type=str, help='old word')
    parser.add_argument('new', type=str, help='new word')
    parser.add_argument('files', type=str, nargs='*', help='docx file path')
    parser.add_argument(
        '-n',
        '--dryrun',
        help='DO NOT save docx file just print replacement result.',
        action='store_true',
        default=False,
    )
    parser.add_argument(
        '-v',
        '--verbose',
        help='print replacement result to stdout',
        action='store_true',
        default=False,
    )
    parser.add_argument('-V', '--version', action='store_true')
    return parser.parse_args()


if __name__ == '__main__':
    # perse()内だとPositional Argument として
    # old, newを要求するからversion表示がされないので、
    # ここであえてversion表示スクリプトをparse()前に走らせる
    if '-V' in sys.argv or '--version' in sys.argv:
        print('replacex:', VERSION)
        sys.exit(0)
    argv = parse()
    main(argv.old,
         argv.new,
         *argv.files,
         dryrun=argv.dryrun,
         verbose=argv.verbose)
